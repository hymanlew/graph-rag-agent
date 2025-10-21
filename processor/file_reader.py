"""
文件读取模块

该模块实现了一个通用的文件读取器，支持多种文件格式的解析和内容提取。
作为知识库构建的基础组件，它负责将各种格式的文档转换为统一的文本格式，
为后续的文本处理和向量生成提供数据基础。
"""
import codecs
import os
from typing import List, Tuple, Dict, Optional
import PyPDF2
from docx import Document
import csv
import json
import yaml
from yaml import CLoader as Loader

from config.settings import FILES_DIR


class FileReader:
    """
    文件读取器类
    
    该类提供统一的接口来读取和解析多种文件格式，是文档处理流程的第一步。
    设计思路是通过工厂模式根据文件扩展名调用不同的解析方法，
    支持递归读取目录，提供统一的错误处理机制。
    
    支持的文件格式：
    - TXT (文本文件)
    - PDF (PDF文档)
    - MD (Markdown文件)
    - DOCX (Word文档)
    - DOC (旧版Word文档)
    - CSV (CSV文件)
    - JSON (JSON文件)
    - YAML/YML (YAML文件)
    """

    def __init__(self, directory_path: str):
        """
        初始化文件读取器
        
        参数：
            directory_path: 文件目录路径，指定要读取的文件所在的根目录
            
        实现思路：
        - 保存目录路径供后续文件读取使用
        - 不执行实际的文件操作，延迟到调用具体方法时才进行
        """
        self.directory_path = directory_path
        
    def read_files(self, file_extensions: Optional[List[str]] = None, recursive: bool = True) -> List[Tuple[str, str]]:
        """
        读取指定扩展名的文件
        
        参数：
            file_extensions: 文件扩展名列表，如 ['.txt', '.pdf']，如不指定则读取所有支持的格式
            recursive: 是否递归读取子目录，默认为True
            
        返回：
            List[Tuple[str, str]]: 文件名和内容的元组列表
            
        实现思路：
        1. 维护支持的文件扩展名与对应处理函数的映射关系
        2. 根据recursive参数决定是否递归读取子目录
        3. 统一的错误处理机制，确保单个文件读取失败不影响整体流程
        4. 提供详细的日志输出，便于调试和监控
        
        业务意义：
        - 作为文档处理的入口函数，提供灵活的文件选择机制
        - 支持多种文件格式，满足不同场景的需求
        - 统一的返回格式，简化上层应用的处理逻辑
        """
        # 文件扩展名和对应的读取方法映射表，实现工厂模式
        supported_extensions = {
            '.txt': self._read_txt,
            '.pdf': self._read_pdf,
            '.md': self._read_markdown,
            '.docx': self._read_docx,
            '.doc': self._read_doc,
            '.csv': self._read_csv,
            '.json': self._read_json,
            '.yaml': self._read_yaml,
            '.yml': self._read_yaml,
        }
        
        # 如未指定扩展名，则使用所有支持的扩展名
        if file_extensions is None:
            file_extensions = list(supported_extensions.keys())
            
        results = []
        try:
            if recursive:
                # 递归读取所有文件
                results = self._read_files_recursive(self.directory_path, file_extensions, supported_extensions)
                print(f"递归读取目录完成，总共读取了 {len(results)} 个文件")
            else:
                # 仅读取当前目录的文件
                all_filenames = os.listdir(self.directory_path)
                print(f"当前目录中共有 {len(all_filenames)} 个文件")
                
                results = self._process_files_in_dir(self.directory_path, all_filenames, file_extensions, supported_extensions)
                print(f"总共读取了 {len(results)} 个文件")
        except Exception as e:
            print(f"列出目录 {self.directory_path} 中的文件时出错: {str(e)}")
            
        return results
    
    def _read_files_recursive(self, root_dir: str, file_extensions: List[str], supported_extensions: Dict) -> List[Tuple[str, str]]:
        """
        递归读取目录及其子目录中的文件
        
        参数：
            root_dir: 当前处理的目录路径
            file_extensions: 要处理的文件扩展名列表
            supported_extensions: 支持的文件扩展名及对应处理函数
            
        返回：
            List[Tuple[str, str]]: 文件名和内容的元组列表
            
        实现思路：
        1. 深度优先搜索(DFS)遍历目录结构
        2. 对于每个子目录，递归调用自身进行处理
        3. 对于文件，检查扩展名并调用对应处理函数
        4. 保存相对路径而非绝对路径，便于后续处理和追踪
        5. 提供详细的日志输出，包括处理进度和错误信息
        """
        results = []
        
        try:
            # 遍历目录内容
            for item in os.listdir(root_dir):
                item_path = os.path.join(root_dir, item)
                
                # 如果是目录，递归处理
                if os.path.isdir(item_path):
                    print(f"递归进入子目录: {item_path}")
                    sub_results = self._read_files_recursive(item_path, file_extensions, supported_extensions)
                    results.extend(sub_results)
                
                # 如果是文件，处理文件
                elif os.path.isfile(item_path):
                    file_ext = os.path.splitext(item)[1].lower()
                    
                    if file_ext in file_extensions:
                        # 获取相对于根目录的路径
                        rel_path = os.path.relpath(item_path, self.directory_path)
                        
                        print(f"处理文件: {rel_path} (类型: {file_ext})")
                        
                        # 使用对应的读取方法处理文件
                        if file_ext in supported_extensions:
                            try:
                                content = supported_extensions[file_ext](item_path)
                                # 存储相对路径而不是仅文件名，以便区分不同目录中的同名文件
                                results.append((rel_path, content))
                                print(f"成功读取文件: {rel_path}, 内容长度: {len(content)}")
                            except Exception as e:
                                print(f"读取文件 {rel_path} 时出错: {str(e)}")
        except Exception as e:
            print(f"列出目录 {root_dir} 中的文件时出错: {str(e)}")
            
        return results
    
    def _process_files_in_dir(self, directory: str, filenames: List[str], file_extensions: List[str], 
                              supported_extensions: Dict) -> List[Tuple[str, str]]:
        """
        处理指定目录中的文件（不递归）
        
        参数：
            directory: 目录路径
            filenames: 文件名列表
            file_extensions: 要处理的文件扩展名列表
            supported_extensions: 支持的文件扩展名及对应处理函数
            
        返回：
            List[Tuple[str, str]]: 文件名和内容的元组列表
            
        实现思路：
        1. 遍历目录中的所有文件名
        2. 检查文件扩展名是否在指定的处理列表中
        3. 使用对应的处理函数读取文件内容
        4. 添加详细的日志输出，记录处理过程
        5. 单个文件处理失败不影响整体流程
        """
        results = []
        
        for filename in filenames:
            file_ext = os.path.splitext(filename)[1].lower()
            
            if file_ext in file_extensions:
                file_path = os.path.join(directory, filename)
                print(f"处理文件: {filename} (类型: {file_ext})")
                
                # 使用对应的读取方法处理文件
                if file_ext in supported_extensions:
                    try:
                        content = supported_extensions[file_ext](file_path)
                        results.append((filename, content))
                        print(f"成功读取文件: {filename}, 内容长度: {len(content)}")
                    except Exception as e:
                        print(f"读取文件 {filename} 时出错: {str(e)}")
        
        return results
    
    def _read_txt(self, file_path: str) -> str:
        """
        读取TXT文件内容
        
        参数：
            file_path: 文件路径
            
        返回：
            str: 文件内容，读取失败时返回错误信息
            
        实现思路：
        1. 首先尝试使用UTF-8编码读取文件
        2. 如果失败，采用编码检测策略：
           - 读取文件前10KB进行编码检测
           - 优先使用chardet库进行智能编码检测
           - 回退到GBK编码（常见于中文文档）
        3. 使用errors='replace'参数确保即使有编码错误也能读取内容
        4. 多层次的异常处理确保函数的鲁棒性
        5. 返回统一格式的错误信息，便于上层应用处理
        """
        try:
            # 首先尝试使用UTF-8编码读取
            with codecs.open(file_path, 'r', encoding='utf-8', errors='replace') as file:
                content = file.read()
            return content
        except Exception as e:
            print(f"读取TXT文件 {os.path.basename(file_path)} 失败: {str(e)}")
            # 尝试使用其他编码
            try:
                with open(file_path, 'rb') as f:
                    raw_data = f.read(10240)  # 读取前10KB进行编码检测
                    try:
                        import chardet
                        result = chardet.detect(raw_data)
                        encoding = result['encoding'] if result['encoding'] else 'gbk'
                    except:
                        encoding = 'gbk'  # 如果chardet不可用，默认使用gbk
                        
                with codecs.open(file_path, 'r', encoding=encoding, errors='replace') as file:
                    content = file.read()
                return content
            except Exception as e2:
                print(f"尝试使用其他编码读取失败: {str(e2)}")
                return f"[无法读取文件内容: {str(e)}]"
            
    def _read_pdf(self, file_path: str) -> str:
        """
        读取PDF文件内容
        
        参数：
            file_path: 文件路径
            
        返回：
            str: 提取的文本内容，读取失败时返回错误信息
            
        实现思路：
        1. 使用PyPDF2库打开PDF文件
        2. 逐页提取文本内容
        3. 添加页与页之间的分隔符（两个换行）
        4. 实现两级异常处理：
           - 页面级异常处理：单页读取失败不影响整体读取
           - 文件级异常处理：整个文件读取失败返回错误信息
        5. 对于无法读取的页面，添加明确的错误标记
        6. 使用or ""确保即使提取结果为None也能正确处理
        """
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num in range(len(pdf_reader.pages)):
                    try:
                        page = pdf_reader.pages[page_num]
                        page_text = page.extract_text() or ""
                        text += page_text + "\n\n"
                    except Exception as e:
                        print(f"读取PDF文件 {os.path.basename(file_path)} 的第 {page_num+1} 页失败: {str(e)}")
                        text += f"[第 {page_num+1} 页无法读取]\n\n"
            return text
        except Exception as e:
            print(f"读取PDF文件 {os.path.basename(file_path)} 失败: {str(e)}")
            return f"[无法读取PDF文件内容: {str(e)}]"
    
    def _read_markdown(self, file_path: str) -> str:
        """
        读取Markdown文件内容
        
        参数：
            file_path: 文件路径
            
        返回：
            str: 文件内容，保留原始Markdown格式，读取失败时返回错误信息
            
        实现思路：
        - 作为纯文本文件处理，保留原始Markdown格式
        - 使用codecs.open确保编码处理正确
        - 设置errors='replace'提高鲁棒性
        - 捕获并处理可能的异常
        
        业务意义：
        - 保留Markdown格式对于后续处理和显示都有帮助
        - 很多技术文档和知识库内容都采用Markdown格式
        """
        try:
            with codecs.open(file_path, 'r', encoding='utf-8', errors='replace') as file:
                md_content = file.read()
                return md_content
        except Exception as e:
            print(f"读取Markdown文件 {os.path.basename(file_path)} 失败: {str(e)}")
            return f"[无法读取Markdown文件内容: {str(e)}]"
    
    def _read_docx(self, file_path: str) -> str:
        """
        读取Word文档(.docx)内容
        
        参数：
            file_path: 文件路径
            
        返回：
            str: 提取的文本内容，读取失败时返回错误信息
            
        实现思路：
        1. 使用python-docx库打开docx文件
        2. 遍历文档中的所有段落对象
        3. 提取每个段落的文本内容
        4. 使用换行符连接所有段落文本
        5. 异常处理确保函数的鲁棒性
        
        业务意义：
        - Word文档是最常见的办公文档格式之一
        - 保留段落结构有助于保持文档的基本逻辑
        """
        try:
            doc = Document(file_path)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            return '\n'.join(full_text)
        except Exception as e:
            print(f"读取Word文档(.docx) {os.path.basename(file_path)} 失败: {str(e)}")
            return f"[无法读取Word文档内容: {str(e)}]"
            
    def _read_doc(self, file_path: str) -> str:
        """
        读取旧版Word文档(.doc)内容
        
        参数：
            file_path: 文件路径
            
        返回：
            str: 提取的文本内容，读取失败时返回错误信息
            
        实现思路：
        采用三级回退策略，逐步尝试不同的方法：
        1. 首选方法：使用Windows系统的win32com接口调用Word应用程序
           - 优点：提取效果最佳，能保留原始格式和内容
           - 缺点：仅支持Windows系统
        2. 次选方法：使用textract库（跨平台）
           - 优点：跨平台支持
           - 缺点：需要安装额外依赖
        3. 备用方法：尝试使用python-docx（不完全兼容）
           - 优点：实现简单，不需要额外安装
           - 缺点：兼容性差，只能部分读取某些.doc文件
        4. 最终回退：返回警告信息，建议转换文件格式
        
        性能优化：
        - 每种方法都检查提取结果是否有效（非空且有实际内容）
        - 提供详细的日志输出，便于调试
        - 合理处理各种可能的异常情况
        
        业务意义：
        - 支持旧版.doc格式，提高系统的兼容性和适用性
        - 多级回退策略确保在不同环境下都有最佳的尝试
        """
        content = ""
        
        # 方法1: 尝试使用win32com (仅Windows)
        try:
            import win32com.client
            
            print(f"尝试使用win32com读取.doc文件: {os.path.basename(file_path)}")
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            
            doc_abs_path = os.path.abspath(file_path)
            doc = word.Documents.Open(doc_abs_path)
            content = doc.Content.Text
            doc.Close()
            word.Quit()
            
            if content and content.strip():
                print(f"使用win32com成功读取.doc文件")
                return content
        except ImportError:
            print("win32com不可用，这不是Windows系统")
        except Exception as e:
            print(f"使用win32com读取.doc失败: {str(e)}")
        
        # 方法2: 尝试使用textract (跨平台)
        try:
            import textract
            print(f"尝试使用textract读取.doc文件: {os.path.basename(file_path)}")
            content = textract.process(file_path).decode('utf-8')
            
            if content and content.strip():
                print(f"使用textract成功读取.doc文件")
                return content
        except ImportError:
            print("textract不可用，请安装: pip install textract")
        except Exception as e:
            print(f"使用textract读取.doc失败: {str(e)}")
        
        # 方法3: 尝试使用python-docx (不完全兼容.doc，但有时可以部分读取)
        try:
            from docx import Document
            print(f"尝试使用python-docx读取.doc文件: {os.path.basename(file_path)}")
            doc = Document(file_path)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            content = '\n'.join(full_text)
            
            if content and content.strip():
                print(f"使用python-docx部分读取.doc文件成功")
                return content
        except ImportError:
            print("python-docx不可用，请安装: pip install python-docx")
        except Exception as e:
            print(f"尝试使用python-docx读取.doc失败: {str(e)}")
        
        # 所有方法都失败，返回警告信息
        warning_msg = f"[警告: 无法读取.doc文件 {os.path.basename(file_path)}，请安装相关依赖或转换为.docx格式]"
        print(warning_msg)
        return warning_msg
    
    def _read_csv(self, file_path: str) -> str:
        """
        读取CSV文件并转换为文本格式
        
        参数：
            file_path: 文件路径
            
        返回：
            str: 转换后的文本内容，读取失败时返回错误信息
            
        实现思路：
        1. 首先尝试使用UTF-8编码读取CSV文件
        2. 使用csv.reader解析CSV格式
        3. 将每一行数据重新用逗号连接成字符串
        4. 使用换行符连接所有行
        5. 编码检测策略与TXT文件类似：
           - 尝试使用chardet进行智能编码检测
           - 回退到GBK编码
        6. 异常处理确保函数的鲁棒性
        
        注意事项：
        - 此方法将CSV转为纯文本，不保留原始的结构化数据
        - 对于需要结构化处理的场景，应使用read_csv_as_dicts方法
        """
        try:
            text = []
            with open(file_path, 'r', encoding='utf-8', errors='replace') as file:
                csv_reader = csv.reader(file)
                for row in csv_reader:
                    text.append(','.join(row))
            return '\n'.join(text)
        except Exception as e:
            print(f"读取CSV文件 {os.path.basename(file_path)} 失败: {str(e)}")
            # 尝试其他编码
            try:
                with open(file_path, 'rb') as f:
                    try:
                        import chardet
                        raw_data = f.read(10240)
                        result = chardet.detect(raw_data)
                        encoding = result['encoding'] if result['encoding'] else 'gbk'
                    except:
                        encoding = 'gbk'  # 如果chardet不可用，默认使用gbk
                        
                text = []
                with open(file_path, 'r', encoding=encoding, errors='replace') as file:
                    csv_reader = csv.reader(file)
                    for row in csv_reader:
                        text.append(','.join(row))
                return '\n'.join(text)
            except Exception as e2:
                print(f"尝试使用其他编码读取CSV失败: {str(e2)}")
                return f"[无法读取CSV文件内容: {str(e)}]"
    
    def read_csv_as_dicts(self, file_path: str) -> List[Dict]:
        """
        读取CSV文件并返回字典列表
        
        参数：
            file_path: 文件路径
            
        返回：
            List[Dict]: CSV数据的字典列表，每一行为一个字典，键为CSV表头
            
        实现思路：
        1. 使用csv.DictReader自动将CSV表头作为字典键
        2. 逐行读取CSV数据并转换为字典
        3. 将每个字典添加到结果列表中
        4. 异常处理确保返回空列表而不是抛出异常
        
        业务意义：
        - 提供结构化的数据访问方式
        - 适合需要按字段处理CSV数据的场景
        - 保留数据的原始结构信息
        """
        try:
            results = []
            with open(file_path, 'r', encoding='utf-8', errors='replace') as file:
                csv_reader = csv.DictReader(file)
                for row in csv_reader:
                    results.append(dict(row))
            return results
        except Exception as e:
            print(f"读取CSV文件为字典列表时出错: {str(e)}")
            return []
    
    def _read_json(self, file_path: str) -> str:
        """
        读取JSON文件并返回格式化的文本
        
        参数：
            file_path: 文件路径
            
        返回：
            str: 格式化的JSON字符串，读取失败时返回错误信息
            
        实现思路：
        1. 读取JSON文件并解析为Python对象
        2. 将对象重新序列化为格式化的JSON字符串
        3. 设置ensure_ascii=False确保非ASCII字符（如中文）正确显示
        4. 设置indent=2使输出更易读
        5. 异常处理确保函数的鲁棒性
        
        设计亮点：
        - 先解析再格式化，确保返回的是有效的JSON格式
        - 美化输出格式，提高可读性
        - 正确处理非ASCII字符
        """
        try:
            with open(file_path, 'r', encoding='utf-8', errors='replace') as file:
                # 加载为对象然后再转为格式化的字符串，以便更好地处理和显示
                data = json.load(file)
                return json.dumps(data, ensure_ascii=False, indent=2)
        except Exception as e:
            print(f"读取JSON文件 {os.path.basename(file_path)} 失败: {str(e)}")
            return f"[无法读取JSON文件内容: {str(e)}]"
    
    def read_json_as_dict(self, file_path: str) -> Dict:
        """
        读取JSON文件并返回字典/列表对象
        
        参数：
            file_path: 文件路径
            
        返回：
            Dict/List: JSON数据对象，可能是字典或列表，取决于JSON结构
            读取失败时返回空字典
            
        实现思路：
        1. 直接使用json.load解析JSON文件
        2. 保留原始数据结构
        3. 异常处理确保返回空字典而不是抛出异常
        
        业务意义：
        - 提供结构化的数据访问方式
        - 适合需要按字段处理JSON数据的场景
        - 保留数据的原始嵌套结构
        """
        try:
            with open(file_path, 'r', encoding='utf-8', errors='replace') as file:
                return json.load(file)
        except Exception as e:
            print(f"读取JSON文件为字典时出错: {str(e)}")
            return {}
    
    def _read_yaml(self, file_path: str) -> str:
        """
        读取YAML文件并返回格式化的文本
        
        参数：
            file_path: 文件路径
            
        返回：
            str: 格式化的YAML字符串，读取失败时返回错误信息
            
        实现思路：
        1. 使用yaml.load和CLoader快速解析YAML文件
        2. 使用yaml.dump重新序列化为格式化的YAML字符串
        3. 设置allow_unicode=True确保非ASCII字符正确显示
        4. 设置default_flow_style=False使用块格式而非流式格式
        5. 异常处理确保函数的鲁棒性
        
        设计亮点：
        - 使用CLoader而非默认Loader以提高解析速度
        - 重新格式化确保输出格式一致和美观
        - 处理非ASCII字符的显示问题
        """
        try:
            with open(file_path, 'r', encoding='utf-8', errors='replace') as file:
                data = yaml.load(file, Loader=Loader)
                # 重新序列化为格式化的YAML字符串
                return yaml.dump(data, allow_unicode=True, default_flow_style=False)
        except Exception as e:
            print(f"读取YAML文件 {os.path.basename(file_path)} 失败: {str(e)}")
            return f"[无法读取YAML文件内容: {str(e)}]"
    
    def read_yaml_as_dict(self, file_path: str) -> Dict:
        """
        读取YAML文件并返回字典对象
        
        参数：
            file_path: 文件路径
            
        返回：
            Dict: YAML数据对象，读取失败时返回空字典
            
        实现思路：
        1. 使用yaml.load和CLoader快速解析YAML文件
        2. 保留原始数据结构
        3. 异常处理确保返回空字典而不是抛出异常
        
        业务意义：
        - 提供结构化的数据访问方式
        - 适合需要按字段处理YAML配置数据的场景
        - 保留数据的原始嵌套结构
        """
        try:
            with open(file_path, 'r', encoding='utf-8', errors='replace') as file:
                return yaml.load(file, Loader=Loader)
        except Exception as e:
            print(f"读取YAML文件为字典时出错: {str(e)}")
            return {}
    
    def read_txt_files(self) -> List[Tuple[str, str]]:
        """
        便捷方法：仅读取所有txt文件
        
        返回：
            List[Tuple[str, str]]: 文件名和内容的元组列表
            
        实现思路：
        - 调用read_files方法并指定仅读取.txt扩展名的文件
        - 提供便捷的接口，简化常见操作
        """
        return self.read_files(['.txt'])
    
    def list_all_files(self, recursive: bool = True) -> List[str]:
        """
        列出目录中的所有文件
        
        参数：
            recursive: 是否递归列出子目录中的文件，默认为True
            
        返回：
            List[str]: 文件路径列表（相对于根目录）
            
        实现思路：
        1. 根据recursive参数决定遍历方式
        2. 递归模式下使用os.walk遍历所有子目录
        3. 非递归模式下直接使用os.listdir
        4. 递归模式下保存相对路径，便于后续处理
        5. 异常处理确保函数的鲁棒性
        
        业务意义：
        - 提供目录内容概览
        - 便于用户了解可处理的文件数量和类型
        - 为后续批量处理提供文件列表
        """
        files = []
        
        try:
            if recursive:
                # 递归遍历所有子目录
                for root, _, filenames in os.walk(self.directory_path):
                    for filename in filenames:
                        # 获取相对于根目录的路径
                        rel_path = os.path.relpath(os.path.join(root, filename), self.directory_path)
                        files.append(rel_path)
            else:
                # 只列出当前目录下的文件
                files = os.listdir(self.directory_path)
        except Exception as e:
            print(f"列出目录文件时出错: {str(e)}")
            
        return files


# 测试代码
if __name__ == '__main__':
    """
    测试函数：验证FileReader类的功能
    
    测试流程：
    1. 初始化文件读取器
    2. 列出目录中的所有文件
    3. 读取所有支持的文件
    4. 统计不同类型文件的数量
    
    目的：
    - 验证文件读取功能是否正常工作
    - 确认支持的各种文件格式能否正确读取
    - 提供使用示例
    """
    print(f"FILES_DIR: {FILES_DIR}")
    # 初始化文件读取器实例
    reader = FileReader(FILES_DIR)
    
    # 列出目录中的所有文件
    print("\n列出目录中的所有文件:")
    all_filenames = reader.list_all_files()
    print(f"目录中共有 {len(all_filenames)} 个文件:")
    for filename in all_filenames:
        print(f"  {filename}")
    
    # 测试读取所有支持的文件
    print("\n读取所有支持的文件:")
    all_files = reader.read_files()
    print(f"成功读取 {len(all_files)} 个文件")
    
    # 显示每种类型文件的数量
    print("\n按类型统计文件:")
    file_types = {}
    for file_name, _ in all_files:
        ext = os.path.splitext(file_name)[1].lower()
        file_types[ext] = file_types.get(ext, 0) + 1
    
    print("Files by type:")
    for ext, count in file_types.items():
        print(f"  {ext}: {count}")